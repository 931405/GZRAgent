"""
Word 模板解析器
解析 NSFC 申请书 Word 模板，提取：
1. 章节结构（标题层级、顺序）
2. 每章节的填充区域（占位文字）
3. 文档样式信息（便于仿制格式）

章节识别策略：
- 优先调用 LLM 对模板中所有标题做整批语义识别（一次接口调用）
- LLM 调用失败时自动退回轻量关键词兜底逻辑
- 识别结果按模板文件路径 + mtime 做内存缓存，避免重复调用
"""

import hashlib
import json
import os
import re
from typing import Dict, List, Optional

# ============================================================
# 标准章节定义（供 LLM prompt 使用，不再用于硬匹配）
# ============================================================

# 标准章节的语义描述，传给 LLM 帮助它做准确判断
STANDARD_SECTIONS: Dict[str, str] = {
    "立项依据": "阐述研究背景、国内外现状、科学问题及本项目的切入点，通常排在最前面",
    "研究目标与内容": "明确总体目标、分解研究内容，以及拟解决的关键科学问题",
    "研究方案与可行性": "描述技术路线、实验/计算方案、可行性分析，可能含'创新性'子节",
    "特色与创新": "单独列出项目的特色与创新之处，也可能叫'创新点'",
    "研究基础": "介绍团队已有的工作基础、研究条件、申请人简历及代表作",
    "参考文献": "文献列表，英文也可能写作 References",
}

# 章节显示顺序
SECTION_ORDER = list(STANDARD_SECTIONS.keys())

# LRU-style 内存缓存：{cache_key: {heading_text: standard_name_or_null}}
_SECTION_MAP_CACHE: Dict[str, Dict[str, Optional[str]]] = {}


# ============================================================
# LLM 整批识别（核心）
# ============================================================

def _build_cache_key(template_path: str) -> str:
    """用文件路径 + 最后修改时间生成缓存 key，文件改变后自动失效。"""
    mtime = str(os.path.getmtime(template_path)) if os.path.exists(template_path) else "0"
    return hashlib.md5(f"{template_path}:{mtime}".encode()).hexdigest()


def _llm_identify_sections(heading_texts: List[str]) -> Dict[str, Optional[str]]:
    """
    将模板中所有标题文本整批发给 LLM，让其判断每个标题对应哪个标准章节。

    Returns:
        {original_heading: standard_section_name}
        无法归属的标题映射到 None。
    """
    if not heading_texts:
        return {}

    standard_list = "\n".join(
        f'  "{name}"：{desc}' for name, desc in STANDARD_SECTIONS.items()
    )
    headings_json = json.dumps(heading_texts, ensure_ascii=False)

    prompt = f"""你是一位熟悉国家自然科学基金（NSFC）申请书格式的专家。
我有一份申请书 Word 模板，其中包含以下标题文字：
{headings_json}

请将每个标题归类到下列标准章节之一（如果某个标题属于子标题或无关内容，请归类为 null）：
{standard_list}

输出要求：
- 只输出一个合法的 JSON 对象，key 是原始标题文字，value 是标准章节名（字符串）或 null
- 不要输出任何解释、注释或代码块标记
- 示例格式：{{"一、立项依据与研究现状": "立项依据", "（二）研究目标": "研究目标与内容", "封面": null}}
"""

    try:
        from src.llm import get_llm
        llm = get_llm(temperature=0.0)
        response = llm.invoke(prompt)
        raw = response.content if hasattr(response, "content") else str(response)

        # 提取 JSON 对象
        match = re.search(r"\{[\s\S]+\}", raw)
        if not match:
            raise ValueError("LLM 输出中未找到 JSON 对象")

        result: dict = json.loads(match.group())
        # 校验 value 只能是 STANDARD_SECTIONS 的 key 或 None
        cleaned: Dict[str, Optional[str]] = {}
        for heading, mapped in result.items():
            if mapped in STANDARD_SECTIONS:
                cleaned[heading] = mapped
            else:
                cleaned[heading] = None
        return cleaned

    except Exception as e:
        print(f"[TemplateParser] LLM 章节识别失败，将使用关键词兜底: {e}")
        return {}


# ============================================================
# 关键词兜底（LLM 失败时使用）
# ============================================================

# 极简关键词表，仅作 fallback
_FALLBACK_KEYWORDS: Dict[str, List[str]] = {
    "立项依据":       ["立项依据", "研究背景", "背景与意义"],
    "研究目标与内容": ["研究目标", "研究内容", "关键科学问题"],
    "研究方案与可行性":["研究方案", "技术路线", "可行性"],
    "特色与创新":     ["特色与创新", "创新点", "创新性"],
    "研究基础":       ["研究基础", "工作基础", "工作条件"],
    "参考文献":       ["参考文献", "references"],
}


def _fallback_normalize(heading_text: str) -> Optional[str]:
    """关键词兜底识别，不区分大小写，仅在 LLM 不可用时使用。"""
    text_lower = heading_text.strip().lower().replace('\u3000', ' ')
    for standard_name, keywords in _FALLBACK_KEYWORDS.items():
        if any(kw.lower() in text_lower for kw in keywords):
            return standard_name
    return None


# ============================================================
# 对外接口：normalize_section_name（带缓存，供 exporter 调用）
# ============================================================

# 运行时全局映射表（由 parse_word_template 构建后注入）
_ACTIVE_SECTION_MAP: Dict[str, Optional[str]] = {}


def normalize_section_name(heading_text: str) -> Optional[str]:
    """
    将标题文本映射到标准章节名。
    优先查询当前已解析模板的 LLM 映射表；未命中时使用轻量关键词兜底。
    """
    text_clean = heading_text.strip()
    # 1. 查精确映射（LLM 结果）
    if text_clean in _ACTIVE_SECTION_MAP:
        return _ACTIVE_SECTION_MAP[text_clean]
    # 2. 关键词兜底
    return _fallback_normalize(text_clean)


def _get_heading_level(style_name: str) -> int:
    """从样式名提取标题级别，默认返回 1。"""
    match = re.search(r'(\d)', style_name)
    if match:
        level = int(match.group(1))
        return min(level, 9)
    cn_map = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5}
    for ch, lvl in cn_map.items():
        if ch in style_name:
            return lvl
    return 1


def parse_word_template(template_path: str, use_llm: bool = True) -> Dict:
    """
    解析 Word 模板文件，提取章节结构。

    流程：
    1. 扫描全部段落，提取样式为"标题"的段落文字
    2. 整批交给 LLM 做语义识别（一次调用），结果按文件缓存
    3. LLM 失败时自动退回关键词兜底
    4. 按识别结果组装 sections 列表

    Args:
        template_path: docx 文件路径
        use_llm      : False 时强制跳过 LLM，只用关键词

    Returns:
        dict:
            has_template      : bool
            sections          : list of {name, original_heading, heading_level,
                                         heading_para_idx, placeholder_text}
            raw_paragraphs    : list of {idx, style, text}
            unmatched_headings: list（未归入任何标准章节的标题）
            section_map       : {original_heading: standard_name} (调试用)
    """
    global _ACTIVE_SECTION_MAP

    try:
        from docx import Document
    except ImportError:
        return _empty_result()

    if not os.path.exists(template_path):
        return _empty_result()

    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"[TemplateParser] 无法打开模板: {e}")
        return _empty_result()

    # ---- 第一遍：收集所有段落 & 标题文字 ----
    raw_paragraphs: List[Dict] = []
    heading_entries: List[Dict] = []   # {idx, style, text, level}

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else "Normal"
        raw_paragraphs.append({"idx": idx, "style": style_name, "text": text})

        is_heading = (
            "Heading" in style_name
            or style_name.startswith("标题")
            or style_name.startswith("heading")
        )
        if is_heading and text:
            heading_entries.append({
                "idx": idx,
                "style": style_name,
                "text": text,
                "level": _get_heading_level(style_name),
            })

    if not heading_entries:
        return _empty_result(raw_paragraphs)

    # ---- 第二遍：整批 LLM 识别（带缓存）----
    heading_texts = [h["text"] for h in heading_entries]
    cache_key = _build_cache_key(template_path)

    if cache_key in _SECTION_MAP_CACHE:
        section_map = _SECTION_MAP_CACHE[cache_key]
        print(f"[TemplateParser] 使用缓存映射（{len(section_map)} 条）")
    else:
        section_map: Dict[str, Optional[str]] = {}
        if use_llm:
            section_map = _llm_identify_sections(heading_texts)

        # LLM 未覆盖的标题用关键词兜底补全
        for text in heading_texts:
            if text not in section_map:
                section_map[text] = _fallback_normalize(text)

        _SECTION_MAP_CACHE[cache_key] = section_map
        print(f"[TemplateParser] LLM 识别完成: "
              f"{sum(1 for v in section_map.values() if v)} / {len(section_map)} 个标题被归类")

    # 注入全局映射，供 normalize_section_name() 快速查询
    _ACTIVE_SECTION_MAP.update(section_map)

    # ---- 第三遍：按映射结果组装 sections ----
    sections: List[Dict] = []
    unmatched_headings: List[Dict] = []
    current_section: Optional[Dict] = None
    content_buffer: List[str] = []

    # 构建 para_idx → heading_entry 的快速查找
    heading_idx_map = {h["idx"]: h for h in heading_entries}

    for row in raw_paragraphs:
        idx, text, style_name = row["idx"], row["text"], row["style"]

        if idx in heading_idx_map and text:
            # 保存上一个章节
            if current_section is not None:
                current_section["placeholder_text"] = "\n".join(content_buffer).strip()
                sections.append(current_section)
                content_buffer = []

            standard_name = section_map.get(text)
            if standard_name:
                current_section = {
                    "name": standard_name,
                    "original_heading": text,
                    "heading_level": heading_idx_map[idx]["level"],
                    "heading_para_idx": idx,
                    "placeholder_text": "",
                }
            else:
                unmatched_headings.append({"idx": idx, "style": style_name, "text": text})
                # 未识别的标题作为上一章节的子内容
                if current_section is not None:
                    content_buffer.append(text)
        else:
            if current_section is not None and text:
                content_buffer.append(text)

    # 末尾章节
    if current_section is not None:
        current_section["placeholder_text"] = "\n".join(content_buffer).strip()
        sections.append(current_section)

    return {
        "has_template": len(sections) > 0,
        "sections": sections,
        "raw_paragraphs": raw_paragraphs,
        "unmatched_headings": unmatched_headings,
        "section_map": section_map,
    }


def _empty_result(raw_paragraphs: Optional[List] = None) -> Dict:
    return {
        "has_template": False,
        "sections": [],
        "raw_paragraphs": raw_paragraphs or [],
        "unmatched_headings": [],
        "section_map": {},
    }


def get_section_order_from_template(template_path: str, use_llm: bool = True) -> List[str]:
    """返回模板中章节的排列顺序（标准名称列表）。"""
    result = parse_word_template(template_path, use_llm=use_llm)
    return [s["name"] for s in result["sections"]]


def invalidate_template_cache(template_path: Optional[str] = None):
    """
    清除 LLM 识别缓存。
    - template_path 为 None 时清除全部缓存
    - 指定路径时只清除对应模板的缓存（模板文件更新后调用）
    """
    global _SECTION_MAP_CACHE, _ACTIVE_SECTION_MAP
    if template_path is None:
        _SECTION_MAP_CACHE.clear()
        _ACTIVE_SECTION_MAP.clear()
        print("[TemplateParser] 全部缓存已清除")
    else:
        key = _build_cache_key(template_path)
        _SECTION_MAP_CACHE.pop(key, None)
        print(f"[TemplateParser] 已清除缓存: {os.path.basename(template_path)}")


def find_template_in_dir(templates_dir: str = "templates") -> Optional[str]:
    """
    在 templates/ 目录下查找最新的 .docx 模板文件。
    Returns: 绝对路径 or None
    """
    if not os.path.exists(templates_dir):
        return None
    docx_files = [
        f for f in os.listdir(templates_dir)
        if f.endswith(".docx") and not f.startswith("~")
    ]
    if not docx_files:
        return None
    # 按修改时间排序，优先使用最新的
    docx_files.sort(
        key=lambda f: os.path.getmtime(os.path.join(templates_dir, f)),
        reverse=True
    )
    return os.path.abspath(os.path.join(templates_dir, docx_files[0]))


def get_template_style_info(template_path: str) -> Dict:
    """
    提取模板的样式参数（字体、字号、页边距等），
    便于生成风格一致的新文档。
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        return {}

    if not os.path.exists(template_path):
        return {}

    info = {}
    try:
        doc = Document(template_path)

        # 页面设置
        if doc.sections:
            sec = doc.sections[0]
            info["page_width"] = sec.page_width
            info["page_height"] = sec.page_height
            info["left_margin"] = sec.left_margin
            info["right_margin"] = sec.right_margin
            info["top_margin"] = sec.top_margin
            info["bottom_margin"] = sec.bottom_margin

        # Normal 样式字体
        normal_style = doc.styles.get("Normal") or doc.styles.get("正文")
        if normal_style and hasattr(normal_style, "font"):
            info["normal_font_name"] = normal_style.font.name
            info["normal_font_size"] = (
                normal_style.font.size.pt if normal_style.font.size else 12
            )

    except Exception as e:
        print(f"[TemplateParser] 提取样式信息失败: {e}")

    return info
