"""
NSFC 申请书导出器（完善版）
===========================
功能：
1. 基于 Word 模板填充内容（识别章节标题，替换占位文字）
2. 无模板时生成标准格式文档（带封面、合适字体、层级标题）
3. 自动插入 Mermaid / SiliconFlow 配图到对应章节
4. 处理 Markdown 格式文本（**加粗**、标题、列表等）
5. 生成前可选是否调用大模型生成配图（generate_images 参数）
"""

import os
import re
import time
from typing import Dict, List, Optional, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ============================================================
# 章节排列顺序（标准）
# ============================================================
SECTION_ORDER = [
    "立项依据",
    "研究目标与内容",
    "研究方案与可行性",
    "特色与创新",
    "研究基础",
    "参考文献",
]


# ============================================================
# 主入口
# ============================================================

def export_to_word(
    project_type: str,
    topic: str,
    drafts: Dict[str, str],
    output_dir: str = "data",
    template_path: Optional[str] = None,
    images: Optional[Dict[str, str]] = None,
    generate_images: bool = True,
    model_provider: Optional[str] = None,
) -> str:
    """
    生成 Word 文档并保存，返回文件绝对路径。

    Args:
        project_type   : 项目类型（面上项目 / 青年基金 …）
        topic          : 研究主题
        drafts         : {section_name: markdown_content}
        output_dir     : 输出目录
        template_path  : 可选的 Word 模板路径；None 时自动查找 templates/ 目录
        images         : 已有配图 {section_name: local_path}；None 时若 generate_images=True 则自动生成
        generate_images: 是否调用大模型生成配图
        model_provider : LLM provider（None = 全局默认）

    Returns:
        str: 保存的 docx 文件绝对路径
    """
    os.makedirs(output_dir, exist_ok=True)

    # ---- 自动查找模板 ----
    if template_path is None:
        from src.utils.word_template_parser import find_template_in_dir
        template_path = find_template_in_dir("templates")

    # ---- 自动生成配图 ----
    if images is None and generate_images and drafts:
        try:
            from src.utils.image_generator import generate_all_diagrams
            print("[Exporter] 正在生成章节配图...")
            images = generate_all_diagrams(drafts, topic, model_provider)
        except Exception as e:
            print(f"[Exporter] 配图生成失败（将继续无图模式）: {e}")
            images = {}

    images = images or {}

    # ---- 选择导出模式 ----
    if template_path and os.path.exists(template_path):
        print(f"[Exporter] 使用模板导出: {template_path}")
        return _export_with_template(
            project_type, topic, drafts, output_dir, template_path, images
        )
    else:
        print("[Exporter] 无模板，生成标准格式文档...")
        return _export_standard(project_type, topic, drafts, output_dir, images)


# ============================================================
# 模式1：基于 Word 模板填充
# ============================================================

def _export_with_template(
    project_type: str,
    topic: str,
    drafts: Dict[str, str],
    output_dir: str,
    template_path: str,
    images: Dict[str, str],
) -> str:
    """
    打开 Word 模板，识别章节标题，将 AI 生成内容填入对应位置，
    同时在章节内容区域后插入配图。
    """
    from src.utils.word_template_parser import normalize_section_name

    doc = Document(template_path)
    paras = doc.paragraphs

    # ---- 建立标题位置索引 ----
    heading_map: List[Dict] = []
    for i, para in enumerate(paras):
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        is_heading = (
            "Heading" in style_name
            or style_name.startswith("标题")
            or style_name.startswith("heading")
        )
        if is_heading and text:
            normalized = normalize_section_name(text)
            if normalized:
                heading_map.append({
                    "section_name": normalized,
                    "para_idx": i,
                    "para_obj": para,
                })

    if not heading_map:
        print("[Exporter] 模板中未识别到 NSFC 章节，退回标准格式。")
        return _export_standard(project_type, topic, drafts, output_dir, images)

    # ---- 倒序处理章节（避免删除行后索引偏移影响后续操作） ----
    for j in range(len(heading_map) - 1, -1, -1):
        entry = heading_map[j]
        section_name = entry["section_name"]
        head_idx = entry["para_idx"]

        next_head_idx = (
            heading_map[j + 1]["para_idx"] if j + 1 < len(heading_map)
            else len(paras)
        )

        content_paras = list(paras[head_idx + 1: next_head_idx])
        new_content = drafts.get(section_name, "")
        if not new_content:
            continue

        # 优先使用 DocumentDOM 渲染（若存在）
        # 这里仅在后端执行中传入的是 state，为了兼容 API 传入的字典这里做个判断
        # 假设 API 调用尚未修改，我们先做一个向后兼容。
        from typing import Any
        dom_node = None
        # 如果 drafts 是完整的状态字典且有 document_dom，则使用它；否则尝试直接解析文本。
        if isinstance(drafts, dict) and "_dom_cache" in drafts and section_name in drafts["_dom_cache"]:
             dom_node = drafts["_dom_cache"][section_name]

        # 清除占位文字
        for p in content_paras:
            _clear_paragraph_content(p)

        # 在标题段后插入内容 + 配图 + DOM 元素
        anchor = paras[head_idx]
        if dom_node:
            _insert_dom_after(doc, anchor, dom_node, section_name, images)
        else:
            _insert_content_after(doc, anchor, new_content, section_name, images)

        # 移除空壳段落
        for p in content_paras:
            parent = p._p.getparent()
            if parent is not None:
                parent.remove(p._p)

    filename = _make_output_path(output_dir, project_type)
    doc.save(filename)
    print(f"[Exporter] 已保存（模板模式）: {filename}")
    return filename


# ============================================================
# 模式2：标准格式文档生成
# ============================================================

def _export_standard(
    project_type: str,
    topic: str,
    drafts: Dict[str, str],
    output_dir: str,
    images: Dict[str, str],
) -> str:
    """生成标准格式 Word 文档（带封面、合适格式）。"""
    doc = Document()
    _setup_document_styles(doc)
    _add_cover(doc, project_type, topic)

    section_keys = SECTION_ORDER + [k for k in drafts if k not in SECTION_ORDER]
    for section_name in section_keys:
        content = drafts.get(section_name, "")
        if not content:
            continue

        heading = doc.add_heading(section_name, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            run.font.name = "黑体"
            run.font.size = Pt(16)

        # 配图插入章节标题后
        _try_insert_image(doc, images.get(section_name), section_name)

        _write_markdown_content(doc, content)
        doc.add_paragraph()

    filename = _make_output_path(output_dir, project_type)
    doc.save(filename)
    print(f"[Exporter] 已保存（标准模式）: {filename}")
    return filename


# ============================================================
# 段落/内容写入工具
# ============================================================

def _insert_content_after(
    doc: Document,
    anchor_para,
    content: str,
    section_name: str,
    images: Dict[str, str],
):
    """在 anchor_para 之后插入：配图 → 正文内容。"""
    insert_ref = anchor_para

    # 先插入配图
    img_path = images.get(section_name)
    if img_path and os.path.exists(img_path):
        insert_ref = _xml_insert_image_para(doc, insert_ref, img_path, section_name)

    # 再插入正文
    for block in _split_content_to_blocks(content):
        insert_ref = _xml_append_text_para(doc, insert_ref, block["text"])


def _xml_append_text_para(doc, anchor, text: str):
    """在 anchor._p 之后插入新段落，返回段落对象。"""
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)

    if text.strip():
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        new_p.append(r)

    # 查找对应的 Paragraph 对象
    for p in doc.paragraphs:
        if p._p is new_p:
            return p

    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc)


def _xml_insert_image_para(doc, anchor, img_path: str, section_name: str):
    """在 anchor 之后插入图片段落，返回图片段落对象。"""
    try:
        cap_p = doc.add_paragraph(f"图：{section_name}技术路线图")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap_p.runs:
            cap_p.runs[0].font.size = Pt(10)
            cap_p.runs[0].font.italic = True

        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(img_path, width=Inches(5.5))

        # 移动到 anchor 后面
        cap_p._p.getparent().remove(cap_p._p)
        img_p._p.getparent().remove(img_p._p)
        anchor._p.addnext(img_p._p)
        anchor._p.addnext(cap_p._p)

        return img_p
    except Exception as e:
        print(f"[Exporter] 模板模式图片插入失败: {e}")
        return anchor


def _try_insert_image(doc: Document, img_path: Optional[str], section_name: str):
    """（标准模式）在当前位置插入图片 + 图注。"""
    if not img_path or not os.path.exists(img_path):
        return
    try:
        caption_para = doc.add_paragraph(f"图：{section_name}技术路线图")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption_para.runs:
            caption_para.runs[0].font.size = Pt(10)
            caption_para.runs[0].font.italic = True
            caption_para.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(5.5))

        doc.add_paragraph()
    except Exception as e:
        print(f"[Exporter] 图片插入失败: {e}")

def _insert_dom_after(
    doc: Document,
    anchor_para,
    dom_node: Any,
    section_name: str,
    images: Dict[str, str],
):
    """根据 DocumentDOM 插入结构化内容（兼容 Phase 4 V2 架构）。"""
    insert_ref = anchor_para

    # 先插入配图
    img_path = images.get(section_name)
    if img_path and os.path.exists(img_path):
        insert_ref = _xml_insert_image_para(doc, insert_ref, img_path, section_name)
        
    for element in dom_node.elements:
        if element.type == "text":
             # 利用原有的 markdown 解析按块写入
             for block in _split_content_to_blocks(element.content):
                 insert_ref = _xml_append_text_para(doc, insert_ref, block["text"])
        elif element.type == "table":
             # 简单的 Markdown Table 转 Word Table（简易实现）
             insert_ref = _xml_append_text_para(doc, insert_ref, f"【生成数据表】:\n{element.content}")
        elif element.type == "formula":
             insert_ref = _xml_append_text_para(doc, insert_ref, f"【数学公式】: {element.content}")
        elif element.type == "image":
             if os.path.exists(element.content):
                  insert_ref = _xml_insert_image_para(doc, insert_ref, element.content, element.id)


def _write_markdown_content(doc: Document, content: str):
    """将 Markdown 格式章节内容写入 doc。"""
    blocks = re.split(r"\n{2,}", content)
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        if block.startswith("## "):
            h = doc.add_heading(block[3:].strip(), level=2)
            _style_heading(h, Pt(14))
            continue

        if block.startswith("### "):
            h = doc.add_heading(block[4:].strip(), level=3)
            _style_heading(h, Pt(13))
            continue

        lines = block.split("\n")

        if all(re.match(r"^[-*]\s", ln.strip()) for ln in lines if ln.strip()):
            for ln in lines:
                ln = ln.strip()
                if re.match(r"^[-*]\s", ln):
                    p = doc.add_paragraph(style="List Bullet")
                    _add_mixed_runs(p, ln[2:].strip())
            continue

        if all(re.match(r"^\d+\.", ln.strip()) for ln in lines if ln.strip()):
            for ln in lines:
                ln = ln.strip()
                m = re.match(r"^\d+\.\s*(.*)", ln)
                if m:
                    p = doc.add_paragraph(style="List Number")
                    _add_mixed_runs(p, m.group(1).strip())
            continue

        combined = " ".join(lines)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        _add_mixed_runs(p, combined)


def _add_mixed_runs(paragraph, text: str):
    """处理 **加粗** 内联标记。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def _split_content_to_blocks(content: str) -> List[Dict]:
    """将 Markdown 内容拆分为纯文本块（用于模板填充模式）。"""
    blocks = []
    for block in re.split(r"\n{2,}", content):
        block = block.strip()
        if not block:
            continue
        clean = re.sub(r"^#{1,3}\s+", "", block)
        clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
        clean = clean.replace("\n", " ").strip()
        if clean:
            blocks.append({"text": clean, "style": "Normal"})
    return blocks


# ============================================================
# 文档样式 / 封面工具
# ============================================================

def _setup_document_styles(doc: Document):
    """设置标准文档的基础字体样式。"""
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(12)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass


def _style_heading(heading, size: Pt):
    for run in heading.runs:
        run.font.size = size
        run.font.name = "黑体"
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def _add_cover(doc: Document, project_type: str, topic: str):
    """生成封面页。"""
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("国家自然科学基金")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run("申 请 书")
    run2.font.size = Pt(20)
    run2.font.bold = True
    run2.font.name = "黑体"

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [("项目类型", project_type), ("研究主题", topic)]:
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = info_p.add_run(f"{label}：")
        r1.font.size = Pt(14)
        r1.font.bold = True
        r2 = info_p.add_run(value)
        r2.font.size = Pt(14)

    doc.add_paragraph()

    try:
        from datetime import date
        date_str = date.today().strftime("%Y 年 %m 月")
    except Exception:
        date_str = "2026 年"

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(date_str).font.size = Pt(12)

    doc.add_page_break()


# ============================================================
# 工具函数
# ============================================================

def _clear_paragraph_content(para):
    for run in para.runs:
        run.text = ""


def _make_output_path(output_dir: str, project_type: str) -> str:
    ts = int(time.time())
    safe_type = re.sub(r"[^\w\u4e00-\u9fff]", "_", project_type)
    return os.path.abspath(os.path.join(output_dir, f"NSFC_{safe_type}_{ts}.docx"))
